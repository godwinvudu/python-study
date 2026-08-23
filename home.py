import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_code_block(doc, code_text):
    # Create a 1x1 table to act as the code container
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    
    # Apply a light gray background shading to the cell
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), 'F2F2F2') # Light gray hex
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Configure paragraph formatting inside the cell
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    # Add code text with Consolas font
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(34, 34, 34)
    
    # Add spacing after the table
    doc.add_paragraph()

# Initialize the document
doc = docx.Document()
doc.add_heading('Code Submission Report', level=1)

# Example code snippet you want to insert
sample_code = (
    "import socket\n\n"
    "def scan_port(target, port):\n"
    "    try:\n"
    "        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)\n"
    "        s.settimeout(1)\n"
    "        s.connect((target, port))\n"
    "        print(f'Port {port} is OPEN')\n"
    "        s.close()\n"
    "    except:\n"
    "        pass"
)

add_code_block(doc, sample_code)

# Save the document
doc.save('formatted_code_output.docx')
print("Document saved successfully as formatted_code_output.docx")